from docx import Document

doc = Document('d:/DOWNLOADS/proj documents/Project Charter Template.docx')

print("=" * 80)
print("PROJECT CHARTER TEMPLATE CONTENT")
print("=" * 80)

# Extract paragraphs
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Extract tables
print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)

for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    for row in table.rows:
        row_text = " | ".join([cell.text.strip() for cell in row.cells])
        print(row_text)
    print("-" * 80)
