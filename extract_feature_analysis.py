from docx import Document
from docx.table import Table

doc = Document('d:/DOWNLOADS/proj documents/Feature Analysis (Event Booking System).docx')

print("="*80)
print("FEATURE ANALYSIS TEMPLATE")
print("="*80)

for element in doc.element.body:
    if element.tag.endswith('p'):
        # Paragraph
        para = next((p for p in doc.paragraphs if p._element == element), None)
        if para and para.text.strip():
            print(para.text)
    elif element.tag.endswith('tbl'):
        # Table
        table = next((t for t in doc.tables if t._element == element), None)
        if table:
            print("\n[TABLE]")
            for i, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.strip().replace('\n', ' ') for cell in row.cells])
                print(row_text)
            print("[/TABLE]\n")
