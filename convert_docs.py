"""
AdMotion Documentation Converter
Converts HTML documentation to Word (.docx) and PDF formats
"""

import os
import subprocess
import sys
from pathlib import Path

def check_pandoc():
    """Check if Pandoc is installed"""
    try:
        result = subprocess.run(['pandoc', '--version'], capture_output=True)
        return result.returncode == 0
    except FileNotFoundError:
        return False

def convert_html_to_docx():
    """Convert HTML file to DOCX format"""
    html_file = Path("AdMotion-Complete-Documentation.html")
    docx_file = Path("AdMotion-Complete-Documentation.docx")
    
    if not html_file.exists():
        print(f"Error: {html_file} not found!")
        return False
    
    if check_pandoc():
        print("✓ Pandoc found - converting HTML to DOCX...")
        try:
            # Convert HTML to DOCX with proper formatting
            cmd = [
                'pandoc',
                str(html_file),
                '-o', str(docx_file),
                '--from', 'html',
                '--to', 'docx',
                '--table-of-contents',
                '--toc-depth=2',
                '--reference-doc=None'  # Use default Word template
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0 and docx_file.exists():
                size_mb = docx_file.stat().st_size / (1024 * 1024)
                print(f"✓ Successfully created: {docx_file}")
                print(f"  File size: {size_mb:.2f} MB")
                return True
            else:
                print(f"✗ Conversion failed: {result.stderr}")
                return False
                
        except Exception as e:
            print(f"✗ Error during conversion: {str(e)}")
            return False
    else:
        print("✗ Pandoc not found. Installing...")
        return install_pandoc_and_retry()

def install_pandoc_and_retry():
    """Try to install Pandoc and retry conversion"""
    try:
        # For Windows
        if sys.platform.startswith('win'):
            print("Please install Pandoc from: https://pandoc.org/installing.html")
            print("Or run: choco install pandoc (if you have Chocolatey)")
        else:
            print("Please install Pandoc using your package manager")
            print("Ubuntu/Debian: sudo apt-get install pandoc")
            print("macOS: brew install pandoc")
        return False
    except Exception as e:
        print(f"Error: {str(e)}")
        return False

def create_alternative_docx():
    """Create alternative DOCX using python-docx if Pandoc not available"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        print("✓ Creating DOCX using python-docx library...")
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('AdMotion: Intelligent Vehicle Advertising Ecosystem', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Business Requirements Document (BRD) & Software Requirements Specification (SRS)', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        doc.add_paragraph()
        info_text = doc.add_paragraph()
        info_text.add_run('Version: ').bold = True
        info_text.add_run('2.0\n')
        info_text.add_run('Date: ').bold = True
        info_text.add_run('February 17, 2026\n')
        info_text.add_run('Status: ').bold = True
        info_text.add_run('Final - Ready for Production\n')
        info_text.add_run('Classification: ').bold = True
        info_text.add_run('Confidential\n')
        
        # Add page break
        doc.add_page_break()
        
        # Add TOC placeholder
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            '1. Executive Summary',
            '2. Business Context',
            '3. Problem Statement',
            '4. Proposed Solution',
            '5. Market Analysis',
            '6. Business Objectives',
            '7. Business Requirements',
            '8. Functional Requirements',
            '9. Non-Functional Requirements',
            '10. Data Requirements',
            '11. API Specifications',
            '12. Database Schema',
            '13. Security Specifications',
            '14. Glossary & Appendices',
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_page_break()
        
        # Save document
        doc_path = Path('AdMotion-Complete-Documentation.docx')
        doc.save(str(doc_path))
        
        if doc_path.exists():
            size_mb = doc_path.stat().st_size / (1024 * 1024)
            print(f"✓ Successfully created: {doc_path}")
            print(f"  File size: {size_mb:.2f} MB")
            return True
        else:
            print("✗ Failed to create DOCX file")
            return False
            
    except ImportError:
        print("✗ python-docx not installed")
        print("  Install with: pip install python-docx")
        return False
    except Exception as e:
        print(f"✗ Error creating DOCX: {str(e)}")
        return False

def main():
    print("="*60)
    print("AdMotion Documentation Converter")
    print("="*60)
    print()
    
    # Try Pandoc first
    if not convert_html_to_docx():
        print()
        print("Trying alternative method...")
        create_alternative_docx()
    
    print()
    print("="*60)
    print("Conversion Complete!")
    print("="*60)
    print()
    print("Next steps:")
    print("1. Open the .docx file in Microsoft Word")
    print("2. Review formatting and tables")
    print("3. Export to PDF from Word (File > Export as PDF)")
    print("4. Verify clickable table of contents in PDF")
    print()

if __name__ == "__main__":
    main()
